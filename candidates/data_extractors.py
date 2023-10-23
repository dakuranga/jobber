import re
import spacy
import PyPDF2
import phonenumbers
from docx import Document

nlp = spacy.load("en_core_web_sm")
email_pattern = r"([a-zA-Z0-9._-]+@[a-zA-Z0-9._-]+\.[a-zA-Z0-9_-]+)"

def extract_phones(text):
    phone_numbers = []
    for match in phonenumbers.PhoneNumberMatcher(text, "US"):
        phone_numbers.append(phonenumbers.format_number(match.number, phonenumbers.PhoneNumberFormat.E164))
    return phone_numbers

def extract_emails(text):
    return re.findall(email_pattern, text)

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text += page.extract_text()
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fullText.append(cell.text)
    return '\n'.join(fullText)

def extract_data_from_cv(file):
    filename = file.name
    if filename.endswith('.pdf'):
        text = extract_text_from_pdf(file)
    elif filename.endswith('.docx'):
        text = extract_text_from_docx(file)
    else:
        return {}

    doc = nlp(text)
    names = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
    name = names[0] if names else ""

    emails = extract_emails(text)
    email = emails[0] if emails else ""

    phones = extract_phones(text)
    phone = phones[0] if phones else ""

    print("Extracted Name:", name)
    print("Extracted Email:", email)
    print("Extracted Phone:", phone)

    return {
        "name": name,
        "email": email,
        "phone": phone
    }
